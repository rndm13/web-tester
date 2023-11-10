from . import model
from docx import Document
# from docx.shared import Cm
from typing import List


def format_float(num: float) -> str:
    if num < 0.001:  # µs
        return f"{round(num * 1000000)}µs"
    if num < 1:  # ms
        return f"{round(num * 1000)}ms"
    return f"{round(num)}s"
        

def export_test_results(filename: str, results: List[model.TestResult]):
    document = Document()

    document.add_heading('Testing Report', 0)

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Endpoint'
    hdr_cells[1].text = 'Severity'
    hdr_cells[2].text = 'Verdict'
    hdr_cells[3].text = 'Response time'
    total_time = 0
    result_count = {
        model.Severity.OK: 0,
        model.Severity.WARNING: 0,
        model.Severity.DANGER: 0,
        model.Severity.CRITICAL: 0,
    }

    for result in results:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{result.endpoint.url} {result.endpoint.http_type()}"
        row_cells[1].text = result.severity.name
        result_count[result.severity] += 1
        row_cells[2].text = result.verdict
        elapsed_seconds = result.elapsed_time.total_seconds()
        row_cells[3].text = format_float(elapsed_seconds)
        total_time += elapsed_seconds
    document.add_paragraph(f"Total response time for tests: {format_float(total_time)}")
    document.add_paragraph(f"Result count (Ok/Warning/Danger/Critical): {result_count[model.Severity.OK]}/{result_count[model.Severity.WARNING]}/{result_count[model.Severity.DANGER]}/{result_count[model.Severity.CRITICAL]}")

    document.save(filename)
